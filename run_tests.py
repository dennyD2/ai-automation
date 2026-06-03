# run_tests.py — Generic AI-Agentic Test Runner
# Works for ANY website, ANY test case sheet, ANY module.
# AI reads each test case, looks at the live page, decides actions one-by-one,
# executes them, observes the result, and loops — no hardcoded steps ever.
# Supports flow document (DOCX) for step-by-step reference.
# Production monitoring with Discord alerts on any failure.

import asyncio
import json
import os
import re
import sys
from typing import Any, Dict, List

import pandas as pd
from playwright.async_api import async_playwright, Page

# ── Config ────────────────────────────────────────────────────────────────────

BASE_URL       = os.getenv("BASE_URL",   "https://hire-qa.bling-ai.com/sagility?reqId=REQ-017239&country=%27US%27&location=%27TX%27&source=SOURCE-3-125&profileID=IND007684")
EXCEL_PATH     = os.getenv("EXCEL_PATH", "Trajector Test cases.xlsx")
#FLOW_DOC_PATH  = os.getenv("FLOW_DOC_PATH", "Flow document.docx")   # optional DOCX flow reference
ARTIFACTS      = "artifacts"
MAX_STEPS      = 15        # increased for flow-driven tests
MAX_RETRIES    = 2         # retries for flaky failures
MODEL          = "deepseek-chat"
API_URL        = "https://api.deepseek.com/v1/chat/completions"
API_KEY        = os.getenv("DEEPSEEK_API_KEY", "")
DISCORD_WEBHOOK = os.getenv("DISCORD_WEBHOOK", "")

# ── Scope: control exactly what to run ───────────────────────────────────────
SCOPE: Dict[str, Any] = {
    "Forgot Password": None,
}   # ← edit this to control what runs

# ── Load flow document (DOCX) if present ─────────────────────────────────────

def load_flow_document(path: str) -> str:
    """Extract text from a DOCX flow document. Returns empty string if not found."""
    if not path or not os.path.exists(path):
        return ""
    try:
        from docx import Document
        doc = Document(path)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # Also extract tables (some flow docs use tables for step/action columns)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        content = "\n".join(lines)
        print(f"📄  Flow document loaded: {path} ({len(content)} chars)")
        return content
    except ImportError:
        print("⚠️  python-docx not installed — flow document will be skipped")
        return ""
    except Exception as e:
        print(f"⚠️  Could not load flow document: {e}")
        return ""

FLOW_DOCUMENT_TEXT = load_flow_document(FLOW_DOC_PATH)

# ── AI call (stdlib only, no langchain) ───────────────────────────────────────

async def call_ai(messages: List[Dict]) -> str:
    import urllib.request
    payload = json.dumps({
        "model": MODEL,
        "temperature": 0,
        "max_tokens": 600,
        "messages": messages,
    }).encode()
    req = urllib.request.Request(
        API_URL,
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {API_KEY}",
        },
    )
    with urllib.request.urlopen(req, timeout=40) as resp:
        data = json.loads(resp.read())
    return data["choices"][0]["message"]["content"]


def extract_json(text: str) -> Dict:
    m = re.search(r"\{.*\}", text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group())
        except Exception:
            pass
    return {}

# ── Page snapshot: what the AI sees ──────────────────────────────────────────

async def page_snapshot(page: Page) -> Dict:
    """Returns a structured, AI-readable summary of the live page."""
    try:
        data = await page.evaluate("""() => {
            function labelFor(el) {
                if (el.id) {
                    const lb = document.querySelector('label[for="' + el.id + '"]');
                    if (lb) return lb.innerText.trim().slice(0,60);
                }
                const prev = el.previousElementSibling;
                if (prev && ['LABEL','SPAN','P','DIV'].includes(prev.tagName))
                    return prev.innerText.trim().slice(0,60);
                return '';
            }
            const inputs = Array.from(
                document.querySelectorAll('input:not([type=hidden]),textarea,select')
            ).filter(e => e.offsetParent !== null).map(e => ({
                type:        e.type || e.tagName.toLowerCase(),
                placeholder: (e.placeholder || '').trim(),
                label:       labelFor(e),
                name:        (e.name  || '').trim(),
                id:          (e.id    || '').trim(),
                value: e.type === 'password' ? '(hidden)' : (e.value || '').trim().slice(0,80),
            }));

            const buttons = Array.from(document.querySelectorAll(
                'button,[role=button],input[type=submit],input[type=button]'
            )).filter(e => e.offsetParent !== null)
              .map(e => (e.innerText || e.value || '').trim().replace(/\\s+/g,' '))
              .filter(t => t.length > 0 && t.length < 100);

            const links = Array.from(document.querySelectorAll('a[href]'))
                .filter(e => e.offsetParent !== null)
                .map(e => ({ text: e.innerText.trim().replace(/\\s+/g,' ').slice(0,60), href: e.href }))
                .filter(l => l.text.length > 0)
                .slice(0, 20);

            const alerts = Array.from(document.querySelectorAll(
                '[role=alert],[class*=error i],[class*=toast i],[class*=message i],[class*=notification i],[class*=snack i],[class*=warning i]'
            )).filter(e => e.offsetParent !== null)
              .map(e => e.innerText.trim().replace(/\\s+/g,' ').slice(0, 300))
              .filter(t => t.length > 0);

            // ── Production health signals ──────────────────────────────────
            const spinners = Array.from(document.querySelectorAll(
                '[class*="spinner" i],[class*="loader" i],[class*="loading" i],' +
                '[aria-label*="loading" i],[class*="skeleton" i],[class*="progress" i]'
            )).filter(e => e.offsetParent !== null).length;

            const isBlank = document.body.innerText.trim().length < 30;

            // Detect bot/chat widget present but no response text
            const botWidget = Array.from(document.querySelectorAll(
                '[class*="chat" i],[class*="bot" i],[class*="widget" i],[id*="chat" i]'
            )).filter(e => e.offsetParent !== null).length > 0;

            return {
                inputs, buttons, links, alerts, botWidget,
                bodyText: document.body.innerText.replace(/\\s+/g,' ').trim().slice(0, 2500),
                spinners,
                isBlank,
            };
        }""")
        data["url"]   = page.url
        data["title"] = await page.title()
        return data
    except Exception as e:
        return {
            "url": page.url, "title": "", "inputs": [], "buttons": [],
            "links": [], "alerts": [], "bodyText": "", "error": str(e),
            "spinners": 0, "isBlank": True, "botWidget": False,
        }


def snapshot_to_text(snap: Dict) -> str:
    parts = [
        f"URL   : {snap.get('url','')}",
        f"Title : {snap.get('title','')}",
    ]
    # Production health flags — surfaced prominently so AI can classify
    if snap.get("isBlank"):
        parts.append("🚨 BLANK PAGE DETECTED — body text is empty or near-empty")
    if snap.get("spinners", 0) > 0:
        parts.append(f"⏳ INFINITE LOADER DETECTED — {snap['spinners']} spinner/loader element(s) visible")
    if snap.get("botWidget"):
        parts.append("🤖 BOT/CHAT WIDGET present on page")

    if snap.get("alerts"):
        parts.append("⚠️  Alerts/errors visible on page:")
        for a in snap["alerts"]:
            parts.append(f"    \"{a}\"")
    if snap.get("inputs"):
        parts.append("📝 Input fields available:")
        for inp in snap["inputs"]:
            desc = inp.get("placeholder") or inp.get("label") or inp.get("name") or inp.get("id") or inp.get("type","?")
            parts.append(f"    [{inp.get('type','input')}] placeholder/label=\"{desc}\"  current-value=\"{inp.get('value','')}\"")
    if snap.get("buttons"):
        parts.append("🔘 Buttons: " + " | ".join(f'"{b}"' for b in snap["buttons"][:15]))
    if snap.get("links"):
        parts.append("🔗 Links : " + " | ".join(f'"{l["text"]}"' for l in snap["links"][:15]))
    parts.append(f"\n📄 Full page text (truncated to 1500 chars):\n{snap.get('bodyText','')[:1500]}")
    return "\n".join(parts)

# ── Execute one AI action ─────────────────────────────────────────────────────

async def execute_action(page: Page, action: Dict) -> str:
    act = action.get("action", "").lower()
    el  = str(action.get("element") or action.get("target") or "").strip()
    val = str(action.get("value", ""))

    # FILL ─────────────────────────────────────────────────────────────────────
    if act == "fill":
        candidates = []
        if el:
            safe = re.escape(el)
            candidates = [
                page.get_by_placeholder(re.compile(safe, re.I)),
                page.get_by_label(re.compile(safe, re.I)),
                page.locator(f'input[name*="{el}" i],textarea[name*="{el}" i]'),
                page.locator(f'input[id*="{el}" i]'),
                page.get_by_role("textbox", name=re.compile(safe, re.I)),
            ]
        candidates.append(page.locator("input:visible,textarea:visible").first)
        for loc in candidates:
            try:
                if await loc.count() > 0:
                    await loc.first.fill(val)
                    try:
                        await loc.first.press("Tab")
                    except Exception:
                        pass
                    return f"Filled \"{el}\" → \"{val[:50]}\""
            except Exception:
                continue
        return f"WARN: could not find input \"{el}\""

    # CLICK ────────────────────────────────────────────────────────────────────
    elif act == "click":
        candidates = []
        if el:
            safe = re.escape(el)
            candidates = [
                page.get_by_role("button", name=re.compile(safe, re.I)),
                page.get_by_role("link",   name=re.compile(safe, re.I)),
                page.get_by_text(re.compile(safe, re.I)),
                page.locator(f'[aria-label*="{el}" i],[title*="{el}" i],[placeholder*="{el}" i]'),
                page.locator(f'[class*="eye"],[class*="toggle"],[class*="show-pass"]'),
            ]
        for loc in candidates:
            try:
                if await loc.count() > 0:
                    await loc.first.click()
                    await page.wait_for_timeout(600)
                    return f"Clicked \"{el}\""
            except Exception:
                continue
        return f"WARN: could not find clickable \"{el}\""

    # PRESS KEY ────────────────────────────────────────────────────────────────
    elif act == "press_key":
        key = action.get("key", "Enter")
        if el:
            safe = re.escape(el)
            for loc in [
                page.get_by_placeholder(re.compile(safe, re.I)),
                page.get_by_label(re.compile(safe, re.I)),
            ]:
                try:
                    if await loc.count() > 0:
                        await loc.first.press(key)
                        await page.wait_for_timeout(400)
                        return f"Pressed {key} on \"{el}\""
                except Exception:
                    continue
        await page.keyboard.press(key)
        await page.wait_for_timeout(400)
        return f"Pressed {key} globally"

    # NAVIGATE ─────────────────────────────────────────────────────────────────
    elif act == "navigate":
        url = action.get("url", BASE_URL)
        try:
            await page.goto(url, wait_until="domcontentloaded", timeout=25000)
            await page.wait_for_function("document.body.innerText.trim().length > 30", timeout=8000)
        except Exception as e:
            return f"WARN: navigate failed — {e}"
        await page.wait_for_timeout(700)
        return f"Navigated to {url}"

    # WAIT ─────────────────────────────────────────────────────────────────────
    elif act == "wait":
        ms = int(action.get("ms", 1000))
        await page.wait_for_timeout(ms)
        return f"Waited {ms}ms"

    # WAIT_FOR_TEXT ────────────────────────────────────────────────────────────
    elif act == "wait_for_text":
        text = action.get("text", "")
        timeout_ms = int(action.get("timeout_ms", 10000))
        try:
            await page.wait_for_selector(f"text={text}", timeout=timeout_ms)
            return f"Text \"{text}\" appeared"
        except Exception:
            return f"WARN: text \"{text}\" did not appear within {timeout_ms}ms — possible timeout/bot no-response"

    else:
        return f"Unknown action \"{act}\" — skipped"

# ── System prompt ─────────────────────────────────────────────────────────────

def build_system_prompt() -> str:
    flow_section = ""
    if FLOW_DOCUMENT_TEXT:
        flow_section = f"""
FLOW DOCUMENT REFERENCE:
The following is the official candidate hiring portal flow. Use it to understand the expected
sequence of steps, screens, and content. Cross-reference it when verifying outcomes.

--- FLOW DOCUMENT START ---
{FLOW_DOCUMENT_TEXT[:4000]}
--- FLOW DOCUMENT END ---

"""

    return f"""You are an autonomous web test automation agent monitoring a production candidate hiring portal.
You execute test cases on a live website by taking one browser action at a time, observing the result,
and deciding the next action.
{flow_section}
You will receive:
1. The test case ID, description, and expected outcome
2. The CURRENT live page state: URL, visible inputs, buttons, links, alerts, page text
3. Production health signals: blank page, infinite loaders, spinners, bot widget status
4. A log of actions already taken in this test

You must respond with EXACTLY ONE JSON object — no prose, no markdown, just the raw JSON.

Available actions:
  {{"action":"fill",          "element":"<placeholder or label text>",      "value":"<text>"}}
  {{"action":"click",         "element":"<button text, link text, or label>"}}
  {{"action":"press_key",     "element":"<field placeholder or label>",     "key":"Tab|Enter|Escape|Backspace"}}
  {{"action":"navigate",      "url":"<full URL>"}}
  {{"action":"wait",          "ms":2000}}
  {{"action":"wait_for_text", "text":"<expected text>",                     "timeout_ms":10000}}
  {{"action":"verify",        "result":"PASS|FAIL|SKIP",                    "reason":"<classified reason>"}}

FAILURE CLASSIFICATION — when calling verify with FAIL, you MUST start the reason with one tag:
  [BLANK_PAGE]       — page body is empty or near-empty (< 30 chars)
  [INFINITE_LOADER]  — spinner/loader visible, page not interactive after waiting
  [BOT_NO_RESPONSE]  — chat/AI bot widget present but no reply after interaction
  [NAV_BROKEN]       — URL did not change or wrong page loaded after navigation
  [API_ERROR]        — alert text mentions network / API / server / 500 / failed
  [ELEMENT_MISSING]  — expected button or field not found on page
  [TIMEOUT]          — page or element did not respond within expected time
  [CONTENT_MISMATCH] — page loaded but expected text or UI state not present

Example: {{"action":"verify","result":"FAIL","reason":"[INFINITE_LOADER] Spinner still visible 10s after clicking Next; page never advanced to Name step."}}

Critical rules:
- Follow the flow document stage order: Application stage first, then Prescreening stage.
- Identify elements by VISIBLE placeholder, label, or button text — never by CSS or ID.
- If a spinner/loader is visible, use wait action (up to 3x) before declaring INFINITE_LOADER.
- If bot widget is present and shows no response after clicking/typing, declare BOT_NO_RESPONSE.
- "Leave field empty" → fill with value "".
- After completing all described steps, always use "verify" to declare PASS or FAIL.
- Verify by comparing what is actually visible on the page to the expected outcome.
- If the test needs real credentials not given, return verify SKIP with a reason.
- Maximum {MAX_STEPS} actions total — you must verify before hitting the limit.
"""

# ── Agentic test loop ─────────────────────────────────────────────────────────

async def run_case(page: Page, case_id: str, description: str, expectation: str) -> Dict:
    result, reason, shot = "FAIL", "", ""
    step_log = []

    system_prompt = build_system_prompt()

    # Navigate to start page with blank-page guard
    try:
        await page.goto(BASE_URL, wait_until="domcontentloaded", timeout=25000)
        await page.wait_for_function("document.body.innerText.trim().length > 30", timeout=8000)
        await page.wait_for_timeout(700)
    except Exception as e:
        err_str = str(e)
        tag = "[TIMEOUT]" if "timeout" in err_str.lower() else "[BLANK_PAGE]"
        shot_path = os.path.join(ARTIFACTS, f"{case_id}_blocked.png")
        try:
            await page.screenshot(path=shot_path, full_page=True)
        except Exception:
            shot_path = ""
        return {
            "case_id": case_id, "description": description,
            "expectation": expectation, "result": "BLOCKED",
            "reason": f"{tag} {err_str}", "steps": [], "screenshot": shot_path,
        }

    conversation = [{"role": "system", "content": system_prompt}]
    history_lines: List[str] = []

    for step_num in range(MAX_STEPS):
        snap  = await page_snapshot(page)
        state = snapshot_to_text(snap)
        hist  = ("\n\nActions taken so far:\n" + "\n".join(history_lines)) if history_lines else ""

        user_content = (
            f"Test ID: {case_id}\n"
            f"Description:\n{description}\n\n"
            f"Expected outcome:\n{expectation}"
            f"{hist}\n\n"
            f"Current page state:\n{state}\n\n"
            f"What is your next action?"
        )

        try:
            ai_raw = await call_ai(conversation + [{"role": "user", "content": user_content}])
        except Exception as e:
            step_log.append({"step": step_num+1, "action": "ai_error", "outcome": str(e)})
            result = "ERROR"
            reason = f"[API_ERROR] AI call failed: {e}"
            break

        action = extract_json(ai_raw)
        act_type = action.get("action", "")
        print(f"      [{step_num+1}] {ai_raw[:150].strip()}")

        # ── verify ────────────────────────────────────────────────────────────
        if act_type == "verify":
            result = action.get("result", "FAIL").upper()
            reason = action.get("reason", "")
            step_log.append({
                "step": step_num+1, "action": "verify",
                "outcome": f"{result}: {reason}",
            })
            break

        # ── execute ───────────────────────────────────────────────────────────
        outcome = await execute_action(page, action)
        await page.wait_for_timeout(500)

        step_log.append({
            "step": step_num+1,
            "action": act_type,
            "element": action.get("element",""),
            "value": action.get("value",""),
            "outcome": outcome,
        })
        history_lines.append(f"  Step {step_num+1}: {act_type}({action.get('element','') or action.get('url','')}) → {outcome}")

        conversation.append({"role": "user",      "content": user_content})
        conversation.append({"role": "assistant", "content": ai_raw})

    else:
        # Forced verify at step limit
        snap  = await page_snapshot(page)
        state = snapshot_to_text(snap)
        forced = (
            f"Step limit reached. Verify now.\n"
            f"Test: {description}\nExpected: {expectation}\n"
            f"Page state:\n{state}"
        )
        try:
            ai_raw = await call_ai(conversation + [{"role": "user", "content": forced}])
            action = extract_json(ai_raw)
            result = action.get("result", "FAIL").upper()
            reason = action.get("reason", "Step limit reached")
        except Exception:
            result = "FAIL"
            reason = "[TIMEOUT] Step limit reached; AI verify failed"

    # Screenshot on non-pass
    if result in ("FAIL", "ERROR", "BLOCKED"):
        os.makedirs(ARTIFACTS, exist_ok=True)
        shot = os.path.join(ARTIFACTS, f"{case_id}.png")
        try:
            await page.screenshot(path=shot, full_page=True)
        except Exception:
            shot = ""
    else:
        shot = ""

    return {
        "case_id":     case_id,
        "description": description,
        "expectation": expectation,
        "result":      result,
        "reason":      reason,
        "steps":       step_log,
        "screenshot":  shot,
    }


# ── Retry wrapper ─────────────────────────────────────────────────────────────

async def run_case_with_retry(page: Page, case_id: str, description: str, expectation: str) -> Dict:
    for attempt in range(1, MAX_RETRIES + 1):
        res = await run_case(page, case_id, description, expectation)
        if res["result"] == "PASS":
            return res
        if attempt < MAX_RETRIES:
            print(f"      🔁  Retry {attempt}/{MAX_RETRIES} for {case_id} (result: {res['result']})")
            await page.wait_for_timeout(3000)
        else:
            return res
    return res  # unreachable but satisfies type checker

# ── Discord alert ─────────────────────────────────────────────────────────────

# Maps failure tags → human-readable label + emoji
FAILURE_TAG_META = {
    "BLANK_PAGE":       ("🌑 Blank Page",        "Website loaded but page is empty"),
    "INFINITE_LOADER":  ("⏳ Infinite Loader",    "Spinner/loader never disappeared"),
    "BOT_NO_RESPONSE":  ("🤖 Bot No Response",    "AI bot widget did not reply"),
    "NAV_BROKEN":       ("🔗 Navigation Broken",  "Page did not navigate as expected"),
    "API_ERROR":        ("🔌 API Error",           "Backend/API call failed"),
    "ELEMENT_MISSING":  ("🔍 Element Missing",     "Expected UI element not found"),
    "TIMEOUT":          ("⏱ Timeout",             "Page or element timed out"),
    "CONTENT_MISMATCH": ("📄 Content Mismatch",   "Page loaded but content is wrong"),
}

async def send_discord_alert(results: List[Dict], run_url: str = ""):
    """Send a rich Discord alert for any non-passing test results."""
    if not DISCORD_WEBHOOK:
        print("  ℹ️  DISCORD_WEBHOOK not set — skipping alert")
        return

    failed = [r for r in results if r["result"] in ("FAIL", "ERROR", "BLOCKED")]
    if not failed:
        print("  ✅  All tests passed — no Discord alert needed")
        return

    total  = len(results)
    passed = sum(1 for r in results if r["result"] == "PASS")

    # Build Discord embed
    fields = []
    for r in failed[:10]:   # Discord embed cap
        reason = r.get("reason", "")
        tag_label = "❓ Unknown"
        tag_detail = ""
        for tag, (label, detail) in FAILURE_TAG_META.items():
            if f"[{tag}]" in reason:
                tag_label  = label
                tag_detail = detail
                break

        fields.append({
            "name": f"{tag_label} — `{r['case_id']}`",
            "value": f"**Reason:** {reason[:180]}\n**Stage:** {r.get('description','')[:100]}",
            "inline": False,
        })

    overflow_note = ""
    if len(failed) > 10:
        overflow_note = f"\n_...and {len(failed) - 10} more failures_"

    embed = {
        "title": f"🚨 Sagility Hiring Portal — Production Monitor Alert",
        "description": (
            f"**{len(failed)}/{total} tests failed** on production run.\n"
            f"✅ Passed: {passed} | ❌ Failed: {len(failed)}\n"
            + (f"[View CI Run]({run_url})" if run_url else "")
            + overflow_note
        ),
        "color": 15158332,   # red
        "fields": fields,
        "footer": {"text": f"Site: {BASE_URL}"},
    }

    import urllib.request
    payload = json.dumps({"embeds": [embed]}).encode()
    req = urllib.request.Request(
        DISCORD_WEBHOOK,
        data=payload,
        headers={"Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=10) as resp:
            print(f"  📣  Discord alert sent — HTTP {resp.status}")
    except Exception as e:
        print(f"  ⚠️  Discord alert failed: {e}")

# ── HTML report ───────────────────────────────────────────────────────────────

def write_html_report(results: List[Dict], path: str):
    colour = {
        "PASS":"#22c55e","FAIL":"#ef4444","SKIP":"#f59e0b",
        "ERROR":"#7c3aed","BLOCKED":"#6b7280",
    }
    rows = ""
    for r in results:
        c = colour.get(r.get("result","FAIL"), "#6b7280")
        steps_html = "<ol>" + "".join(
            f"<li><b>{s.get('action','')}</b>"
            + (f" [{s.get('element','')}]" if s.get('element') else "")
            + (f" = \"{s.get('value','')}\"" if s.get('value') else "")
            + f" → <i>{s.get('outcome','')}</i></li>"
            for s in r.get("steps", [])
        ) + "</ol>"
        shot_html = (
            f'<a href="../{r["screenshot"]}" target="_blank">📷</a>'
            if r.get("screenshot") else ""
        )
        # Highlight failure tag in reason
        reason_display = r.get("reason", "")
        for tag in FAILURE_TAG_META:
            if f"[{tag}]" in reason_display:
                label, _ = FAILURE_TAG_META[tag]
                reason_display = reason_display.replace(
                    f"[{tag}]",
                    f'<span style="background:#fef2f2;color:#b91c1c;padding:1px 5px;border-radius:4px;font-size:11px;font-weight:bold">{label}</span>'
                )
                break

        rows += f"""<tr>
          <td style="white-space:nowrap">{r['case_id']}</td>
          <td style="font-size:12px;white-space:pre-wrap">{r.get('description','')}</td>
          <td style="font-size:12px">{r.get('expectation','')}</td>
          <td style="font-size:12px;color:#444">{reason_display}</td>
          <td style="font-size:11px">{steps_html}</td>
          <td style="font-weight:bold;color:{c};font-size:18px;text-align:center">{r.get('result','')}</td>
          <td style="text-align:center">{shot_html}</td>
        </tr>"""

    total  = len(results)
    passed = sum(1 for r in results if r.get("result")=="PASS")
    failed = sum(1 for r in results if r.get("result")=="FAIL")
    errors = sum(1 for r in results if r.get("result") in ("ERROR","BLOCKED"))
    skipped= sum(1 for r in results if r.get("result")=="SKIP")

    # Failure breakdown by tag
    tag_counts = {}
    for r in results:
        reason = r.get("reason", "")
        for tag in FAILURE_TAG_META:
            if f"[{tag}]" in reason:
                tag_counts[tag] = tag_counts.get(tag, 0) + 1
                break

    tag_badges = ""
    for tag, count in tag_counts.items():
        label, _ = FAILURE_TAG_META[tag]
        tag_badges += f'<div class="badge" style="background:#fef2f2;color:#b91c1c;border:1px solid #fecaca">{label}: {count}</div>'

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<title>AI Test Report — Sagility Production Monitor</title>
<style>
  body{{font-family:Arial,sans-serif;margin:24px;background:#f8fafc;color:#1e293b}}
  h1{{margin-bottom:8px}}
  .meta{{color:#64748b;font-size:13px;margin-bottom:16px}}
  .summary{{display:flex;gap:12px;flex-wrap:wrap;margin-bottom:16px}}
  .badge{{padding:10px 18px;border-radius:8px;color:#fff;font-weight:bold;font-size:14px}}
  .tag-row{{display:flex;gap:8px;flex-wrap:wrap;margin-bottom:24px}}
  table{{border-collapse:collapse;width:100%;background:#fff;border-radius:10px;
         overflow:hidden;box-shadow:0 1px 6px #00000018}}
  th{{background:#1e293b;color:#fff;padding:10px 8px;font-size:13px;text-align:left}}
  td{{padding:8px;border-bottom:1px solid #e2e8f0;vertical-align:top;font-size:13px}}
  tr:hover{{background:#f1f5f9}}
  ol{{margin:4px 0;padding-left:18px}}li{{margin-bottom:2px}}
</style></head><body>
<h1>🤖 AI Agentic Test Report — Sagility Production Monitor</h1>
<p class="meta">File: <b>{os.path.basename(EXCEL_PATH)}</b> &nbsp;|&nbsp; Site: <b>{BASE_URL}</b>
{"&nbsp;|&nbsp; Flow doc: <b>" + os.path.basename(FLOW_DOC_PATH) + "</b>" if FLOW_DOCUMENT_TEXT else ""}</p>
<div class="summary">
  <div class="badge" style="background:#22c55e">✅ PASS: {passed}</div>
  <div class="badge" style="background:#ef4444">❌ FAIL: {failed}</div>
  <div class="badge" style="background:#7c3aed">💥 ERROR/BLOCKED: {errors}</div>
  <div class="badge" style="background:#f59e0b">⏭ SKIP: {skipped}</div>
  <div class="badge" style="background:#64748b">📋 TOTAL: {total}</div>
</div>
{"<div class='tag-row'>" + tag_badges + "</div>" if tag_badges else ""}
<table>
  <tr>
    <th>Case ID</th><th>Description</th><th>Expected</th>
    <th>AI Reasoning</th><th>Steps Taken by AI</th><th>Result</th><th>📷</th>
  </tr>
  {rows}
</table></body></html>"""

    with open(path, "w", encoding="utf-8") as fh:
        fh.write(html)

# ── helpers ───────────────────────────────────────────────────────────────────

def to_str(x: Any) -> str:
    if x is None:
        return ""
    if isinstance(x, float) and pd.isna(x):
        return ""
    return str(x).strip()

# ── entry point ───────────────────────────────────────────────────────────────

async def main():
    if not API_KEY:
        print("❌  DEEPSEEK_API_KEY is not set.")
        print("    Export it:  export DEEPSEEK_API_KEY=sk-...")
        sys.exit(1)

    os.makedirs(ARTIFACTS, exist_ok=True)
    all_results: List[Dict] = []
    out_json = os.path.join(ARTIFACTS, "results.json")
    out_html = os.path.join(ARTIFACTS, "report.html")

    # CI run URL for Discord link (set by GitHub Actions)
    run_url = ""
    gh_server = os.getenv("GITHUB_SERVER_URL", "")
    gh_repo   = os.getenv("GITHUB_REPOSITORY", "")
    gh_run_id = os.getenv("GITHUB_RUN_ID", "")
    if gh_server and gh_repo and gh_run_id:
        run_url = f"{gh_server}/{gh_repo}/actions/runs/{gh_run_id}"

    # Write placeholder so CI always has an artifact to upload
    with open(out_json, "w") as fh:
        json.dump([], fh)

    # Load all sheets
    xl = pd.ExcelFile(EXCEL_PATH)
    print(f"📂  Excel: {EXCEL_PATH}")
    print(f"📋  Sheets: {xl.sheet_names}")
    if FLOW_DOCUMENT_TEXT:
        print(f"📄  Flow document active ({len(FLOW_DOCUMENT_TEXT)} chars injected into AI context)")

    # Resolve which sheets to run
    if SCOPE:
        sheets_to_run = [s for s in xl.sheet_names if s in SCOPE]
        print(f"🎯  Scope active — running sheets: {sheets_to_run}")
    else:
        sheets_to_run = xl.sheet_names
        print(f"🎯  No scope — running ALL sheets")

    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        context = await browser.new_context(viewport={"width": 1280, "height": 800})
        page    = await context.new_page()
        page.set_default_timeout(15000)

        # ── Preflight ─────────────────────────────────────────────────────────
        try:
            await page.goto(BASE_URL, wait_until="domcontentloaded", timeout=25000)
            await page.wait_for_function("document.body.innerText.trim().length > 30", timeout=8000)
            print(f"\n✅  Preflight OK — {BASE_URL}")
        except Exception as exc:
            tag = "[TIMEOUT]" if "timeout" in str(exc).lower() else "[BLANK_PAGE]"
            print(f"\n❌  Preflight FAILED: {exc}")
            preflight_result = {
                "case_id": "preflight", "description": "Connectivity check",
                "expectation": BASE_URL, "result": "BLOCKED",
                "reason": f"{tag} {exc}", "steps": [], "screenshot": "",
            }
            all_results.append(preflight_result)
            with open(out_json, "w") as fh:
                json.dump(all_results, fh, indent=2)
            write_html_report(all_results, out_html)
            await browser.close()
            # Send Discord alert immediately on preflight failure
            await send_discord_alert(all_results, run_url)
            sys.exit(2)

        # ── Run sheets ────────────────────────────────────────────────────────
        for sheet in sheets_to_run:
            print(f"\n{'='*60}")
            print(f"  Sheet: {sheet}")
            print(f"{'='*60}")
            try:
                df = pd.read_excel(EXCEL_PATH, sheet_name=sheet)
            except Exception as exc:
                print(f"  ⚠️  Skipping sheet '{sheet}': {exc}")
                continue

            allowed_ids = SCOPE.get(sheet) if SCOPE else None

            for idx, row in df.iterrows():
                case_id     = to_str(row.get("Test case ID", row.get("Test case", ""))) or f"{sheet}_row{idx}"
                description = to_str(row.get("Description", ""))
                expectation = to_str(row.get("Expectation", ""))

                if not description:
                    continue

                if allowed_ids is not None and case_id not in allowed_ids:
                    continue

                print(f"\n  🚀  {case_id}")
                print(f"      📝  {description[:120].replace(chr(10),' ')}")

                res = await run_case_with_retry(page, case_id, description, expectation)

                emoji = {"PASS":"✅","FAIL":"❌","SKIP":"⏭","ERROR":"💥","BLOCKED":"🚫"}.get(res["result"],"❓")
                print(f"      {emoji}  {res['result']} — {str(res.get('reason',''))[:120]}")

                all_results.append(res)

                # Save after every test so partial results survive a crash
                with open(out_json, "w", encoding="utf-8") as fh:
                    json.dump(all_results, fh, indent=2, ensure_ascii=False)

        await browser.close()

    write_html_report(all_results, out_html)

    total  = len(all_results)
    passed = sum(1 for r in all_results if r.get("result") == "PASS")
    print(f"\n{'='*60}")
    print(f"  📊  {passed}/{total} passed")
    print(f"  📄  HTML report  →  {out_html}")
    print(f"  📄  JSON results →  {out_json}")

    # Send Discord alert for any failures
    await send_discord_alert(all_results, run_url)


if __name__ == "__main__":
    asyncio.run(main())
